from docx import Document


class DOCXLoader:
    def load(self, file_obj) -> str:
        try:
            doc = Document(file_obj)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text_parts.append(row_text)

            return "\n".join(text_parts)
        except Exception as e:
            raise RuntimeError(f"Failed to load DOCX: {str(e)}")
